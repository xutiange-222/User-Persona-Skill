"""
preprocess.py — 访谈文件格式归一化

把用户上传的 docx/txt/xlsx 转成统一的 .txt 文件,放到工作目录的 processed/ 下。

用法:
    python preprocess.py \\
        --input-dir <用户上传目录> \\
        --workdir <工作目录> \\
        --groups '<角色分组 JSON>'

groups 参数格式:
    [
      {"name": "PM", "files": ["path/to/张三.docx", "path/to/李四.docx"]},
      {"name": "设计师", "files": ["path/to/王五.docx"]}
    ]

如果 groups 不传,默认所有文件归到一个 "default" 组。
"""

import argparse
import json
import logging
import os
import re
import shutil
import sys
from pathlib import Path

from path_utils import OUTPUT_ROOT_NAME, resolve_process_dir

logger = logging.getLogger("preprocess")
logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")


def sanitize_filename(name: str) -> str:
    """清洗文件名:去空格、特殊字符,保留中文和基本符号"""
    stem = Path(name).stem
    # 替换空格、斜杠、反斜杠、冒号
    stem = re.sub(r"[\s/\\:*?\"<>|]+", "_", stem)
    return f"{stem}.txt"


def extract_docx(file_path: Path) -> str:
    """提取 .docx 文件正文(去页眉页脚)"""
    try:
        from docx import Document
    except ImportError:
        logger.error("需要安装 python-docx: pip install python-docx")
        raise

    doc = Document(str(file_path))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # 也提取表格内容(访谈记录有时是表格)
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                paragraphs.append(" | ".join(row_texts))

    return "\n".join(paragraphs)


def extract_xlsx(file_path: Path) -> str:
    """提取 .xlsx 文件,转为 markdown 表格(灵活处理三种格式)"""
    try:
        import openpyxl
    except ImportError:
        logger.error("需要安装 openpyxl: pip install openpyxl")
        raise

    wb = openpyxl.load_workbook(str(file_path), data_only=True)
    output = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        if ws.max_row == 0:
            continue

        output.append(f"## Sheet: {sheet_name}\n")

        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            # 跳过完全空行
            if not any(c.strip() for c in cells):
                continue
            rows.append(cells)

        if not rows:
            continue

        # 转 markdown 表格
        max_cols = max(len(r) for r in rows)
        # 表头
        header = rows[0] + [""] * (max_cols - len(rows[0]))
        output.append("| " + " | ".join(header) + " |")
        output.append("| " + " | ".join(["---"] * max_cols) + " |")
        # 数据行
        for row in rows[1:]:
            padded = row + [""] * (max_cols - len(row))
            output.append("| " + " | ".join(padded) + " |")

        output.append("")  # 空行分隔 sheet

    return "\n".join(output)


def extract_txt(file_path: Path) -> str:
    """读 .txt 文件,自动检测编码"""
    encodings = ["utf-8", "gb18030", "gbk", "utf-16"]
    for enc in encodings:
        try:
            with open(file_path, "r", encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise ValueError(f"无法识别 {file_path} 的编码,试过 {encodings}")


def extract_json(file_path: Path) -> str:
    """读 .json 文件,转成可读文本"""
    with open(file_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    # 如果是数组,每条转一段
    if isinstance(data, list):
        return "\n\n".join(json.dumps(item, ensure_ascii=False, indent=2) for item in data)
    return json.dumps(data, ensure_ascii=False, indent=2)


def process_file(file_path: Path, output_path: Path) -> bool:
    """处理单个文件,返回是否成功"""
    suffix = file_path.suffix.lower()
    try:
        if suffix == ".docx":
            content = extract_docx(file_path)
        elif suffix in (".xlsx", ".xls"):
            content = extract_xlsx(file_path)
        elif suffix == ".txt":
            content = extract_txt(file_path)
        elif suffix == ".json":
            content = extract_json(file_path)
        else:
            logger.warning(f"不支持的格式: {file_path.name},跳过")
            return False

        if not content.strip():
            logger.warning(f"文件内容为空: {file_path.name}")
            return False

        output_path.parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)
        logger.info(f"✓ {file_path.name} → {output_path}")
        return True

    except Exception as e:
        logger.error(f"✗ 处理失败 {file_path.name}: {e}")
        return False


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input-dir", help="用户上传目录(可选,如果用 --groups 传完整路径就不需要)")
    parser.add_argument("--workdir", required=True, help=f"项目运行目录或过程稿目录,默认结构位于 {OUTPUT_ROOT_NAME}/<项目名>-<日期时间>/过程稿")
    parser.add_argument("--groups", help="角色分组 JSON 字符串")
    args = parser.parse_args()

    workdir = resolve_process_dir(Path(args.workdir))
    workdir.mkdir(parents=True, exist_ok=True)
    processed_root = workdir / "processed"

    # 解析 groups
    if args.groups:
        groups = json.loads(args.groups)
    elif args.input_dir:
        # 默认:所有文件归到 default 组
        input_dir = Path(args.input_dir)
        files = []
        for ext in (".docx", ".txt", ".xlsx", ".xls", ".json"):
            files.extend(input_dir.glob(f"**/*{ext}"))
        groups = [{"name": "default", "files": [str(f) for f in files]}]
    else:
        logger.error("必须提供 --input-dir 或 --groups")
        sys.exit(1)

    error_log = []
    success_count = 0
    total_count = 0

    for group in groups:
        group_name = group["name"]
        group_dir = processed_root / group_name
        group_dir.mkdir(parents=True, exist_ok=True)

        for file_path_str in group["files"]:
            file_path = Path(file_path_str)
            total_count += 1
            if not file_path.exists():
                logger.error(f"文件不存在: {file_path}")
                error_log.append(f"{file_path}: 不存在")
                continue

            output_path = group_dir / sanitize_filename(file_path.name)
            if process_file(file_path, output_path):
                success_count += 1
            else:
                error_log.append(f"{file_path}: 处理失败")

    # 写 error log
    if error_log:
        with open(workdir / "preprocess_errors.log", "w", encoding="utf-8") as f:
            f.write("\n".join(error_log))

    # 输出汇总
    summary = {
        "total": total_count,
        "success": success_count,
        "failed": total_count - success_count,
        "groups": [{"name": g["name"], "file_count": len(g["files"])} for g in groups],
        "processed_dir": str(processed_root),
    }
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
